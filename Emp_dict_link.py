import csv
from docx import Document
import requests
import json
import re
from bs4 import BeautifulSoup
import io
import zipfile

# Clean text function
def clean_text(text):
    if text is not None:
        if any(tag in text for tag in ['<', '>']):
            clean_text = " ".join(text.split())
            soup = BeautifulSoup(clean_text, 'html.parser')
            return soup.get_text()
        else:
            return text.strip()
    else:
        return ""


# Remove all symbol except this
def remove_symbol(text):
    if text is not None:
        cleaned_text = re.sub('[^a-zA-Zа-яА-ЯәғқңөұүhіӘҒҚҢӨҰҮҺІ0-9,.()-/@:;!№#$%&?*=_«» ]', '', text)
        return cleaned_text
    else:
        return None


def extract_email(data):
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    for item in data:
        if re.match(email_pattern, item):
            return item

    return None




# Reading docx file
def read_docx_tables(file_path):
    doc = Document(io.BytesIO(file_path)) #Временное хранилище
    tables = []
    for table in doc.tables:
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text)
            tables.append(row_data)
    return tables


item_dict={}

url = "https://www.gov.kz/graphql"

languages = ['en', 'kk', 'ru']

# Request of Postman
for l in languages:
    headers = {
        'Accept-Language': l
    }
    body = """
    {
      projectdetails (_size:>0){
       id
        emp_dict{
                full_text{
                   description
                   document
                          }
                }

      }
    }
    """
    r = requests.post(url=url, headers=headers, json={"query": body})
    data = json.loads(r.text)
    organizations = data['data']['projectdetails']

    #Writing  in csv file
    with open("Emp_dict_files.csv", "w", encoding='utf-8') as csv_file:
        writer = csv.writer(csv_file)
        for i in organizations:
            project_id = i['id']

            # Change all None for zero
            if i['emp_dict'] is None:
                i['emp_dict'] ={"full_text":"o"}
            if i['emp_dict']['full_text'] is None:
                i['emp_dict']['full_text'] ={"document":"0"}
            if i['emp_dict']['full_text'] is None:
                i['emp_dict']['full_text'] ={"description":"0"}

            full_text = i['emp_dict']['full_text']

            #Get only not None and remove symbol and concatenate
            if isinstance(full_text, dict):
                if full_text['document'] is not None and full_text['document'] != '0':
                    full_text['document'] = remove_symbol(full_text['document']).strip()
                    full_url = 'https://www.gov.kz/' + full_text['document']

                    item_dict[project_id] = [
                        project_id,
                        full_text['description'],
                        full_url
                    ]
items_arr = list(item_dict.values())

header = ["Project ID", "Description", "Url"]

with open("Emp_dict_link.csv", "w", encoding='utf-8') as f:
    w = csv.writer(f, delimiter=",")
    w.writerow(header)
    w.writerows(items_arr)












#    ВЫГРУЗКА ВСЕХ ДАННЫХ
# import csv
# from docx import Document
# import requests
# import json
# import re
# from bs4 import BeautifulSoup
# import io
# import zipfile
#
# # Clean text function
# def clean_text(text):
#     if text is not None:
#         if any(tag in text for tag in ['<', '>']):
#             clean_text = " ".join(text.split())
#             soup = BeautifulSoup(clean_text, 'html.parser')
#             return soup.get_text()
#         else:
#             return text.strip()
#     else:
#         return ""
#
#
# # Remove all symbol except this
# def remove_symbol(text):
#     if text is not None:
#         cleaned_text = re.sub('[^a-zA-Zа-яА-ЯәғқңөұүhіӘҒҚҢӨҰҮҺІ0-9,.()-/@:;!№#$%&?*=_«» ]', '', text)
#         return cleaned_text
#     else:
#         return None
#
# # Reading docx file
# def read_docx_tables(file_path):
#     doc = Document(io.BytesIO(file_path)) #Временное хранилище
#     tables = []
#     for table in doc.tables:
#         for row in table.rows:
#             row_data = []
#             for cell in row.cells:
#                 row_data.append(cell.text)
#             tables.append(row_data)
#     return tables
#
# url = "https://www.gov.kz/graphql"
#
# languages = ['en', 'kk', 'ru']
#
# # Request of Postman
# for l in languages:
#     headers = {
#         'Accept-Language': l
#     }
#     body = """
#     {
#       projectdetails (_size:>0){
#        id
#         emp_dict{
#                 full_text{
#                    description
#                    document
#                           }
#                 }
#
#       }
#     }
#     """
#     r = requests.post(url=url, headers=headers, json={"query": body})
#     data = json.loads(r.text)
#     organizations = data['data']['projectdetails']
#
#     #Writing  in csv file
#     with open("Emp_dict_document.csv", 'w', newline='', encoding='utf-8') as csvfile:
#         writer = csv.writer(csvfile)
#
#         for i in organizations:
#             project_id = i['id']
#
#             # Change all None for zero
#             if i['emp_dict'] is None:
#                 i['emp_dict'] ={"full_text":"o"}
#             if i['emp_dict']['full_text'] is None:
#                 i['emp_dict']['full_text'] ={"document":"0"}
#             if i['emp_dict']['full_text'] is None:
#                 i['emp_dict']['full_text'] ={"description":"0"}
#
#             full_text = i['emp_dict']['full_text']
#
#             #Get only not None and remove symbol and concatenate
#             if isinstance(full_text, dict):
#                 if full_text['document'] is not None:
#                     full_text['document'] = remove_symbol(full_text['document']).strip()
#                     if full_text['document'].endswith('.docx'):
#                         full_url = 'https://www.gov.kz/' + full_text['document']
#
#
#                         # Download file
#                         query_parameters = {"downloadformat": "csv"}
#                         response = requests.get(full_url, params=query_parameters)
#
#
#                         # Reading docx file and write in csv file
#                         try:
#                             tables = read_docx_tables(response.content)
#                             for row in tables:
#                                 writer.writerow([project_id] + [full_text['description']] + row)
#                         except zipfile.BadZipfile:
#                             print("File is not a valid Docx file:", full_url)








#   ВЫГРУЗКА ОПРЕДЕЛЕННЫЙ ДАННЫХ
# import csv
# from docx import Document
# import requests
# import json
# import re
# from bs4 import BeautifulSoup
# import io
# import zipfile
#
# # Clean text function
# def clean_text(text):
#     if text is not None:
#         if any(tag in text for tag in ['<', '>']):
#             clean_text = " ".join(text.split())
#             soup = BeautifulSoup(clean_text, 'html.parser')
#             return soup.get_text()
#         else:
#             return text.strip()
#     else:
#         return ""
#
#
# # Remove all symbol except this
# def remove_symbol(text):
#     if text is not None:
#         cleaned_text = re.sub('[^a-zA-Zа-яА-ЯәғқңөұүhіӘҒҚҢӨҰҮҺІ0-9,.()-/@:;!№#$%&?*=_«» ]', '', text)
#         return cleaned_text
#     else:
#         return None
#
#
# def extract_email(data):
#     email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
#     for item in data:
#         if re.match(email_pattern, item):
#             return item
#
#     return None
#
#
# def extract_phone_number(data):
#     phone_pattern = r'\b\d{2,}-\d{2,}-\d{2,}\b'
#     for item in data:
#         if re.match(phone_pattern, item):
#             return item
#
#     return None
#
#
#
# def extract_phone_number2(data):
#     phone_pattern = r'\b(?:\+?\d{3,4} ?)?\(?\d{5}\)? ?\d{2,3}-\d{2}-\d{2}\b'
#     for item in data:
#         match = re.search(phone_pattern, item)
#         if match:
#             return match.group()
#     return None
#
#
#
# def extract_phone_number3(data):
#     phone_pattern = r'\b(?:\+?\d{1,3} ?)?\(?\d{5,}\)? ?\d{0,3}-?\d{2}-?\d{2}\b'
#     for item in data:
#         match = re.search(phone_pattern, item)
#         if match:
#             return match.group()
#     return None
#
# def extract_full_name(data):
#     name_pattern = r'([А-ЯA-ZЁӘҒҚӨҰҮІ][а-яa-zёәғқөұүңіғһ]+(?:\s+[А-ЯA-ZЁӘҒҚӨҰҮІ][а-яa-zёәғқөұүңіғһ]+){1,2})'
#
#     for item in data:
#         match = re.match(name_pattern, item.strip())
#         if match:
#             return match.group().strip()
#
#     return None
#
# # Reading docx file
# def read_docx_tables(file_path):
#     doc = Document(io.BytesIO(file_path)) #Временное хранилище
#     tables = []
#     for table in doc.tables:
#         for row in table.rows:
#             row_data = []
#             for cell in row.cells:
#                 row_data.append(cell.text)
#             tables.append(row_data)
#     return tables
#
#
#
#
# url = "https://www.gov.kz/graphql"
#
# languages = ['en', 'kk', 'ru']
#
# item_dict = {}
#
# # Request of Postman
# for l in languages:
#     headers = {
#         'Accept-Language': l
#     }
#     body = """
#     {
#       projectdetails (_size:>0){
#        id
#         emp_dict{
#                 full_text{
#                    description
#                    document
#                           }
#                 }
#
#       }
#     }
#     """
#     r = requests.post(url=url, headers=headers, json={"query": body})
#     data = json.loads(r.text)
#     organizations = data['data']['projectdetails']
#
#     #Writing  in csv file
#     for i in organizations:
#         project_id = i['id']
#
#         # Change all None for zero
#         if i['emp_dict'] is None:
#             i['emp_dict'] ={"full_text":"o"}
#         if i['emp_dict']['full_text'] is None:
#             i['emp_dict']['full_text'] ={"document":"0"}
#         if i['emp_dict']['full_text'] is None:
#             i['emp_dict']['full_text'] ={"description":"0"}
#
#         full_text = i['emp_dict']['full_text']
#
#         #Get only not None and remove symbol and concatenate
#         if isinstance(full_text, dict):
#             if full_text['document'] is not None:
#                 full_text['document'] = remove_symbol(full_text['document']).strip()
#                 if full_text['document'].endswith('.docx'):
#                     full_url = 'https://www.gov.kz/' + full_text['document']
#
#
#                     # Download file
#                     query_parameters = {"downloadformat": "csv"}
#                     response = requests.get(full_url, params=query_parameters)
#
#
#
#
#                     # Reading docx file and write in csv file
#                     my_list = read_docx_tables(response.content)
#
#
#
#                     for row in my_list:
#                         if len(set(row)) >1:
#                             item_dict[project_id]=[
#                                 project_id,
#                                 full_text['description'],
#                                 extract_full_name(row),
#                                 extract_email(row),
#                                 extract_phone_number(row),
#                                 extract_phone_number3(row)
#                              ]
#
# item = list(item_dict.values())
#
# header = ['Project_id', 'Description', 'Full_Name', 'Email', 'Phone_Number',
#           'Phone_Number2']
#
# with open("Emp_dict.csv", "a") as f:
#     w = csv.writer(f, delimiter=",")
#     w.writerow(header)
#     w.writerows(item)
















